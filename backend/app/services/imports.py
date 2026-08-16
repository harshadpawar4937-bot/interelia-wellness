"""CSV / DOCX import helpers for products and CRM customers."""

from __future__ import annotations

import csv
import io
import re
from decimal import Decimal, InvalidOperation
from typing import Any, Dict, Iterable, List, Optional

from sqlalchemy.orm import Session

from app.models import Brand, Category, CrmCustomer, Product

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def _slugify(text: str, fallback: str = "item") -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")
    return (s[:80] or fallback)


def _dec(value: Any, default: Decimal = Decimal("0")) -> Decimal:
    if value is None:
        return default
    s = str(value).strip().replace(",", "")
    if not s:
        return default
    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return default


def _int_qty(value: Any) -> int:
    if value is None:
        return 0
    s = str(value).strip()
    if not s:
        return 0
    if ":" in s:
        s = s.split(":", 1)[0]
    try:
        return max(0, int(float(s)))
    except (ValueError, TypeError):
        return 0


def _norm_key(k: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (k or "").strip().lower())


def _row_map(row: Dict[str, str]) -> Dict[str, str]:
    return {_norm_key(k): (v or "").strip() for k, v in row.items() if k is not None}


def _pick(mapped: Dict[str, str], *aliases: str) -> str:
    for a in aliases:
        v = mapped.get(_norm_key(a))
        if v:
            return v
    return ""


def _boolish(value: str, default: bool = True) -> bool:
    s = (value or "").strip().lower()
    if not s:
        return default
    if s in ("1", "true", "yes", "y", "opted_in", "opt-in", "active"):
        return True
    if s in ("0", "false", "no", "n", "opted_out", "opt-out", "inactive"):
        return False
    return default


def _normalize_phone(raw: str) -> Optional[str]:
    if not raw:
        return None
    digits = re.sub(r"[^\d+]", "", raw)
    if len(re.sub(r"\D", "", digits)) < 7:
        return None
    return digits[:32]


def _normalize_email(raw: str) -> Optional[str]:
    if not raw:
        return None
    email = raw.strip().lower()
    if not EMAIL_RE.match(email):
        return None
    return email[:255]


CATEGORY_SLUG_MAP = {
    "drug": "medicines",
    "otc": "wellness",
    "derma": "personal-care",
    "surgical": "medicines",
    "cosmetic": "personal-care",
    "generic": "medicines",
    "ayurvedic": "ayurveda",
    "nutraceuticals": "nutrition",
    "dpco": "medicines",
    "coldchain": "medicines",
    "sp offer": "wellness",
    "beverages": "nutrition",
}


def ensure_category(db: Session, raw: str, cache: Dict[str, Category]) -> Category:
    key = (raw or "general").strip() or "general"
    mapped = CATEGORY_SLUG_MAP.get(key.lower(), _slugify(key, "general"))
    if mapped in cache:
        return cache[mapped]
    cat = db.query(Category).filter_by(slug=mapped).first()
    if not cat:
        cat = Category(name=key.title() if key != "general" else "General", slug=mapped)
        db.add(cat)
        db.flush()
    cache[mapped] = cat
    return cat


def ensure_brand(db: Session, raw: str, cache: Dict[str, Brand]) -> Brand:
    name = (raw or "Interelia").strip() or "Interelia"
    if name.lower() in ("not available", "na", "n/a", "-"):
        name = "Interelia"
    if name in cache:
        return cache[name]
    brand = db.query(Brand).filter_by(name=name).first()
    if not brand:
        from app.services.serializers import slugify

        base = slugify(name)
        slug = base
        n = 2
        while db.query(Brand).filter_by(slug=slug).first():
            slug = f"{base}-{n}"
            n += 1
        brand = Brand(name=name, slug=slug, is_partner=True, is_active=True)
        db.add(brand)
        db.flush()
    elif not brand.slug:
        from app.services.serializers import slugify

        brand.slug = slugify(brand.name)
    cache[name] = brand
    return brand


def detect_csv_kind(fieldnames: Optional[List[str]]) -> str:
    """Return 'products', 'customers', or 'unknown' from headers."""
    if not fieldnames:
        return "unknown"
    keys = {_norm_key(f) for f in fieldnames}
    product_hits = {"id", "name", "mrp", "ptr", "currentstock", "manufacturer", "packaging"}
    customer_hits = {
        "fullname",
        "email",
        "phone",
        "mobile",
        "customername",
        "customerno",
        "contactno",
        "address",
        "city",
        "company",
        "doctorname",
        "profilename",
        "netoflbilledon",
        "nettotalamount",
    }

    if len(keys & product_hits) >= 3 and "mrp" in keys:
        return "products"
    if len(keys & customer_hits) >= 2:
        return "customers"
    if "suppliername" in keys and "mrp" in keys:
        return "products"
    return "unknown"


def import_itemwise_products(
    db: Session,
    file_text: str,
    *,
    limit: Optional[int] = None,
) -> Dict[str, Any]:
    """Upsert products from ITEMWISE_STOCK_REPORT-style CSV."""
    import hashlib

    reader = csv.DictReader(io.StringIO(file_text))
    kind = detect_csv_kind(reader.fieldnames)
    if kind == "customers":
        return {
            "created": 0,
            "updated": 0,
            "skipped": 0,
            "total": 0,
            "errors": [{"row": 0, "error": "This looks like a customer CSV. Import it on the Customers tab."}],
            "detail": "Wrong file type for Products",
        }

    created = updated = skipped = 0
    errors: List[Dict[str, Any]] = []
    cat_cache: Dict[str, Category] = {}
    brand_cache: Dict[str, Brand] = {}
    pending_slugs: set = set()
    batch = 0

    for i, row in enumerate(reader, start=2):
        if limit is not None and (i - 2) >= limit:
            break
        mapped = _row_map(row)
        sku = _pick(mapped, "Id", "SKU", "sku", "Product Id", "Item Id")
        name = _pick(mapped, "Name", "Product Name", "product", "Item Name")
        if not name:
            skipped += 1
            if len(errors) < 50:
                errors.append({"row": i, "error": "Missing product name"})
            continue

        packaging = _pick(mapped, "Packaging", "Pack", "Pack Size")
        rack = _pick(mapped, "Rack")
        category_raw = _pick(mapped, "Category")
        manufacturer = _pick(mapped, "Manufacturer", "Brand")
        mrp = _dec(_pick(mapped, "MRP"), Decimal("0"))
        ptr = _dec(_pick(mapped, "PTR"))
        stock = _int_qty(_pick(mapped, "Current Stock", "Stock", "stock_qty"))
        strip_qty = _int_qty(_pick(mapped, "Current Strip Qty"))
        loose_qty = _int_qty(_pick(mapped, "Current Loose Qty"))
        b2c_strip = _int_qty(_pick(mapped, "B2C Strip Qty"))
        b2c_loose = _int_qty(_pick(mapped, "B2C Loose Qty"))
        b2c_sale = _int_qty(_pick(mapped, "B2C Sale Qty"))
        b2b_sale = _int_qty(_pick(mapped, "B2B Sale Qty"))
        stk_xfer = _int_qty(_pick(mapped, "Stk Transfer Qty"))
        tot_strip = _int_qty(_pick(mapped, "Total Strip Qty"))
        tot_loose = _int_qty(_pick(mapped, "Total Loose Qty"))
        tot_sale = _int_qty(_pick(mapped, "Total Sale Qty"))
        purchase_qty = _pick(mapped, "Purchase Qty")[:64] or None
        margin = _dec(_pick(mapped, "Purchase Margin %", "Purchase Margin"))
        supplier = _pick(mapped, "Supplier Name", "Supplier")

        if mrp <= 0 and ptr > 0:
            mrp = ptr
        if mrp <= 0:
            mrp = Decimal("1")
        price = ptr if ptr > 0 else mrp
        if price <= 0:
            price = mrp

        stock_fields = dict(
            stock_qty=stock,
            current_strip_qty=strip_qty,
            current_loose_qty=loose_qty,
            b2c_strip_qty=b2c_strip,
            b2c_loose_qty=b2c_loose,
            b2c_sale_qty=b2c_sale,
            b2b_sale_qty=b2b_sale,
            stk_transfer_qty=stk_xfer,
            total_strip_qty=tot_strip,
            total_loose_qty=tot_loose,
            total_sale_qty=tot_sale,
            purchase_qty=purchase_qty,
            purchase_margin_pct=margin,
            pack_size=packaging[:100] if packaging else None,
            rack=rack[:64] if rack else None,
            supplier_name=supplier[:255] if supplier else None,
        )

        try:
            nested = db.begin_nested()
            cat = ensure_category(db, category_raw, cat_cache)
            brand = ensure_brand(db, manufacturer, brand_cache)
            requires_rx = category_raw.lower() == "drug"

            existing: Optional[Product] = None
            if sku:
                existing = db.query(Product).filter_by(sku=sku).first()

            if not existing:
                slug_base = _slugify(name)
                if sku:
                    digest = hashlib.md5(sku.encode("utf-8")).hexdigest()[:10]
                    slug = f"{slug_base}-{digest}"
                else:
                    slug = slug_base
                slug = slug[:240]
                n = 1
                candidate = slug
                while candidate in pending_slugs or db.query(Product).filter_by(slug=candidate).first():
                    n += 1
                    candidate = f"{slug}-{n}"
                pending_slugs.add(candidate)
                db.add(
                    Product(
                        sku=sku or None,
                        name=name[:255],
                        slug=candidate,
                        description=f"{name}. Pack: {packaging}".strip(),
                        price=price,
                        mrp=mrp,
                        requires_prescription=requires_rx,
                        category_id=cat.id,
                        brand_id=brand.id,
                        is_active=True,
                        meta_title=name[:160],
                        **stock_fields,
                    )
                )
                created += 1
            else:
                existing.name = name[:255]
                existing.price = price
                existing.mrp = mrp
                existing.requires_prescription = requires_rx
                existing.category_id = cat.id
                existing.brand_id = brand.id
                existing.is_active = True
                for k, v in stock_fields.items():
                    if v is None and k in ("pack_size", "rack", "supplier_name", "purchase_qty"):
                        continue
                    setattr(existing, k, v)
                updated += 1

            nested.commit()
            batch += 1
            if batch % 200 == 0:
                db.commit()
                pending_slugs.clear()
        except Exception as exc:  # noqa: BLE001
            try:
                nested.rollback()
            except Exception:  # noqa: BLE001
                db.rollback()
            skipped += 1
            if len(errors) < 50:
                errors.append({"row": i, "error": str(exc)[:200]})
            continue

    db.commit()
    return {
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "total": created + updated + skipped,
        "errors": errors,
    }


def export_products_csv(products: Iterable[Product]) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(
        [
            "Id",
            "Name",
            "Packaging",
            "Rack",
            "Category",
            "Manufacturer",
            "MRP",
            "PTR",
            "Current Strip Qty",
            "Current Loose Qty",
            "Current Stock",
            "B2C Strip Qty",
            "B2C Loose Qty",
            "B2C Sale Qty",
            "B2B Sale Qty",
            "Stk Transfer Qty",
            "Total Strip Qty",
            "Total Loose Qty",
            "Total Sale Qty",
            "Purchase Qty",
            "Purchase Margin %",
            "Supplier Name",
        ]
    )
    for p in products:
        writer.writerow(
            [
                p.sku or "",
                p.name,
                p.pack_size or "",
                p.rack or "",
                p.category.name if p.category else "",
                p.brand.name if p.brand else "",
                str(p.mrp),
                str(p.price),
                getattr(p, "current_strip_qty", 0) or 0,
                getattr(p, "current_loose_qty", 0) or 0,
                p.stock_qty,
                getattr(p, "b2c_strip_qty", 0) or 0,
                getattr(p, "b2c_loose_qty", 0) or 0,
                getattr(p, "b2c_sale_qty", 0) or 0,
                getattr(p, "b2b_sale_qty", 0) or 0,
                getattr(p, "stk_transfer_qty", 0) or 0,
                getattr(p, "total_strip_qty", 0) or 0,
                getattr(p, "total_loose_qty", 0) or 0,
                getattr(p, "total_sale_qty", 0) or 0,
                getattr(p, "purchase_qty", None) or "",
                str(getattr(p, "purchase_margin_pct", 0) or 0),
                p.supplier_name or "",
            ]
        )
    return buf.getvalue()


CUSTOMER_TEMPLATE_HEADERS = [
    "Customer No.",
    "Name",
    "Contact No",
    "Discount",
    "Profile Name",
    "Doctor Name",
    "Family Name",
    "Payment Mode",
    "Vouchers",
    "Address",
    "City",
    "No. of Bills",
    "Last Billed On",
    "Net Total Amount",
    "Total Due Amount",
]


def customer_csv_template() -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(CUSTOMER_TEMPLATE_HEADERS)
    writer.writerow(
        [
            "C10001",
            "Priya Sharma",
            "9876543210",
            "0",
            "Acme Green",
            "-",
            "-",
            "-",
            "0",
            "12 MG Road",
            "Mumbai",
            "1",
            "01-May-2026┊05:16 PM",
            "500",
            "0",
        ]
    )
    return buf.getvalue()


def _blankish(value: Any) -> Optional[str]:
    if value is None:
        return None
    s = str(value).strip()
    if not s or s in ("-", "—", "NA", "N/A", "null", "None"):
        return None
    return s


def _rows_from_customer_file(raw: bytes, filename: str = "") -> List[Dict[str, Any]]:
    """Parse CUSTOMER_REPORT as xlsx (often misnamed .csv) or real CSV."""
    name = (filename or "").lower()
    is_zip = raw[:2] == b"PK"

    if is_zip or name.endswith((".xlsx", ".xlsm")) or (name.endswith(".csv") and is_zip):
        import tempfile
        from pathlib import Path
        from openpyxl import load_workbook

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = Path(tmp.name)
        try:
            wb = load_workbook(tmp_path, read_only=True, data_only=True)
            ws = wb[wb.sheetnames[0]]
            it = ws.iter_rows(values_only=True)
            headers = [str(h).strip() if h is not None else "" for h in next(it)]
            out: List[Dict[str, Any]] = []
            for row in it:
                out.append({headers[i]: row[i] if i < len(row) else None for i in range(len(headers))})
            wb.close()
            return out
        finally:
            tmp_path.unlink(missing_ok=True)

    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    text = text.replace("\x00", "")
    reader = csv.DictReader(io.StringIO(text))
    return list(reader)


def import_customers_file(db: Session, raw: bytes, filename: str = "") -> Dict[str, Any]:
    """Import CUSTOMER_REPORT (xlsx/csv) or generic CRM CSV into crm_customers."""
    try:
        rows = _rows_from_customer_file(raw, filename)
    except Exception as e:  # noqa: BLE001
        return {
            "created": 0,
            "updated": 0,
            "skipped": 0,
            "total": 0,
            "errors": [{"row": 0, "error": f"Could not read file: {e}"}],
            "detail": "Unreadable customer file",
        }

    if not rows:
        return {
            "created": 0,
            "updated": 0,
            "skipped": 0,
            "total": 0,
            "errors": [{"row": 0, "error": "No data rows found"}],
            "detail": "Empty file",
        }

    headers = list(rows[0].keys())
    kind = detect_csv_kind(headers)
    if kind == "products":
        # Allow stock CSV → suppliers path via text
        try:
            text = raw.decode("utf-8-sig")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="ignore")
        return import_suppliers_from_stock_csv(db, text)

    created = updated = skipped = 0
    errors: List[Dict[str, Any]] = []
    seen_ext: set = set()
    seen_phones: set = set()
    batch = 0

    for i, row in enumerate(rows, start=2):
        mapped = _row_map({str(k): "" if v is None else str(v) for k, v in row.items()})
        # Prefer pharmacy CUSTOMER_REPORT headers
        external_id = _blankish(
            _pick(mapped, "Customer No.", "Customer No", "Customer Number", "external_id", "Customer Id", "Id", "Code")
        )
        name = _blankish(_pick(mapped, "Name", "full_name", "Full Name", "Customer Name", "Patient Name"))
        phone = _normalize_phone(_pick(mapped, "Contact No", "Contact No.", "phone", "Phone", "Mobile", "WhatsApp") or "")
        email = _normalize_email(_pick(mapped, "email", "Email", "E-mail") or "")
        discount = _dec(_pick(mapped, "Discount", "discount_pct"))
        profile = _blankish(_pick(mapped, "Profile Name", "profile_name"))
        doctor = _blankish(_pick(mapped, "Doctor Name", "doctor_name"))
        family = _blankish(_pick(mapped, "Family Name", "family_name"))
        pay_mode = _blankish(_pick(mapped, "Payment Mode", "payment_mode"))
        vouchers = _int_qty(_pick(mapped, "Vouchers"))
        address = _blankish(_pick(mapped, "Address", "address", "Customer Address"))
        city = _blankish(_pick(mapped, "City", "city"))
        bills = _int_qty(_pick(mapped, "No. of Bills", "bills_count", "Bills"))
        last_billed = _blankish(_pick(mapped, "Last Billed On", "last_billed_on"))
        net_total = _dec(_pick(mapped, "Net Total Amount", "net_total_amount"))
        due_total = _dec(_pick(mapped, "Total Due Amount", "total_due_amount"))
        company = _blankish(_pick(mapped, "company", "Company"))
        state = _blankish(_pick(mapped, "state", "State"))
        pincode = _blankish(_pick(mapped, "pincode", "Pincode", "Pin Code"))
        tags = _blankish(_pick(mapped, "tags", "Tags"))
        notes = _blankish(_pick(mapped, "notes", "Notes"))
        opt_in = _boolish(_pick(mapped, "marketing_opt_in", "Opt In"), False)

        if not name and not phone and not external_id and not email:
            skipped += 1
            continue
        if not name:
            name = external_id or phone or email or "Customer"

        if external_id and external_id in seen_ext:
            skipped += 1
            if len(errors) < 100:
                errors.append({"row": i, "error": f"Duplicate Customer No. in file: {external_id}"})
            continue
        if phone and phone in seen_phones and not external_id:
            skipped += 1
            continue
        if external_id:
            seen_ext.add(external_id)
        if phone:
            seen_phones.add(phone)

        existing = _find_customer(db, external_id=external_id, email=email, phone=phone)
        payload = {
            "full_name": name[:150],
            "email": email,
            "phone": phone,
            "company": company[:150] if company else None,
            "discount_pct": discount,
            "profile_name": profile[:120] if profile else None,
            "doctor_name": doctor[:150] if doctor else None,
            "family_name": family[:150] if family else None,
            "payment_mode": pay_mode[:64] if pay_mode else None,
            "vouchers": vouchers,
            "address": address,
            "city": city[:100] if city else None,
            "state": state[:100] if state else None,
            "pincode": pincode[:20] if pincode else None,
            "bills_count": bills,
            "last_billed_on": last_billed[:64] if last_billed else None,
            "net_total_amount": net_total,
            "total_due_amount": due_total,
            "tags": tags[:500] if tags else None,
            "notes": notes,
            "marketing_opt_in": opt_in,
            "is_active": True,
            "source": "customer_report",
        }
        if external_id:
            payload["external_id"] = external_id[:120]

        if existing:
            for k, v in payload.items():
                if v is None and k in (
                    "email",
                    "phone",
                    "company",
                    "address",
                    "city",
                    "state",
                    "pincode",
                    "tags",
                    "notes",
                    "profile_name",
                    "doctor_name",
                    "family_name",
                    "payment_mode",
                    "last_billed_on",
                ):
                    continue
                setattr(existing, k, v)
            updated += 1
        else:
            db.add(CrmCustomer(**payload))
            created += 1

        batch += 1
        if batch % 300 == 0:
            db.flush()

    db.commit()
    return {
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "total": created + updated + skipped,
        "errors": errors,
        "detail": None,
    }


def import_customers_csv(db: Session, file_text: str) -> Dict[str, Any]:
    """Back-compat wrapper for plain CSV text."""
    return import_customers_file(db, file_text.encode("utf-8"), "customers.csv")


def _find_customer(
    db: Session,
    *,
    external_id: Optional[str],
    email: Optional[str],
    phone: Optional[str],
) -> Optional[CrmCustomer]:
    if external_id:
        c = db.query(CrmCustomer).filter_by(external_id=external_id).first()
        if c:
            return c
    if email:
        c = db.query(CrmCustomer).filter_by(email=email).first()
        if c:
            return c
    if phone:
        c = db.query(CrmCustomer).filter_by(phone=phone).first()
        if c:
            return c
    return None


def import_suppliers_from_stock_csv(db: Session, file_text: str) -> Dict[str, Any]:
    """Pull unique Supplier Name values from ITEMWISE stock CSV into CRM."""
    reader = csv.DictReader(io.StringIO(file_text))
    created = updated = skipped = 0
    errors: List[Dict[str, Any]] = []
    seen: set = set()

    for i, row in enumerate(reader, start=2):
        mapped = _row_map(row)
        supplier = _pick(mapped, "Supplier Name", "Supplier")
        if not supplier or supplier.lower() in ("not available", "na", "n/a", "-"):
            skipped += 1
            continue
        key = supplier.strip().lower()
        if key in seen:
            skipped += 1
            continue
        seen.add(key)

        existing = (
            db.query(CrmCustomer)
            .filter(CrmCustomer.company == supplier[:150])
            .first()
        )
        if not existing:
            existing = (
                db.query(CrmCustomer)
                .filter(CrmCustomer.full_name == supplier[:150])
                .first()
            )

        if existing:
            existing.company = supplier[:150]
            existing.tags = ",".join(
                sorted(set((existing.tags or "").split(",") + ["supplier"]) - {""})
            )
            existing.source = existing.source or "stock_csv"
            existing.is_active = True
            updated += 1
        else:
            db.add(
                CrmCustomer(
                    full_name=supplier[:150],
                    company=supplier[:150],
                    tags="supplier",
                    source="stock_csv",
                    marketing_opt_in=False,
                    is_active=True,
                    notes="Imported from ITEMWISE stock report (Supplier Name)",
                )
            )
            created += 1

        if (created + updated) % 200 == 0:
            db.flush()

    db.commit()
    return {
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "total": created + updated + skipped,
        "errors": errors,
        "detail": f"Imported {len(seen)} unique suppliers from stock report",
    }


def export_customers_csv(customers: Iterable[CrmCustomer]) -> str:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(CUSTOMER_TEMPLATE_HEADERS)
    for c in customers:
        writer.writerow(
            [
                c.external_id or "",
                c.full_name,
                c.phone or "",
                str(getattr(c, "discount_pct", 0) or 0),
                getattr(c, "profile_name", None) or "",
                getattr(c, "doctor_name", None) or "",
                getattr(c, "family_name", None) or "",
                getattr(c, "payment_mode", None) or "",
                getattr(c, "vouchers", 0) or 0,
                c.address or "",
                c.city or "",
                getattr(c, "bills_count", 0) or 0,
                getattr(c, "last_billed_on", None) or "",
                str(getattr(c, "net_total_amount", 0) or 0),
                str(getattr(c, "total_due_amount", 0) or 0),
            ]
        )
    return buf.getvalue()


def import_customer_from_coa_docx(db: Session, file_bytes: bytes) -> Dict[str, Any]:
    """Extract Customer Name / Address from COA Form 39 DOCX and upsert CRM customer."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is required for COA import") from e

    doc = Document(io.BytesIO(file_bytes))
    fields: Dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            for i, cell in enumerate(cells):
                label = cell.rstrip(":").strip().lower()
                if label in ("customer name", "customer address", "sample name", "batch no.", "batch no"):
                    for j in range(i + 1, len(cells)):
                        if cells[j] and cells[j] not in (":",):
                            fields[label] = cells[j]
                            break

    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for idx, line in enumerate(paras):
        key = line.rstrip(":").strip().lower()
        if key in ("customer name", "customer address", "sample name") and idx + 1 < len(paras):
            nxt = paras[idx + 1]
            if nxt not in (":",) and key not in fields:
                fields[key] = nxt

    name = fields.get("customer name", "").strip()
    address = fields.get("customer address", "").strip()
    sample = fields.get("sample name", "").strip()

    if not name or name.upper().startswith("XXX"):
        return {
            "created": 0,
            "updated": 0,
            "skipped": 1,
            "total": 1,
            "errors": [{"row": 0, "error": "No real Customer Name found in COA"}],
            "detail": "No real Customer Name found in COA",
        }

    notes_parts = []
    if sample and not sample.upper().startswith("XXX"):
        notes_parts.append(f"COA sample: {sample}")
    batch = fields.get("batch no.") or fields.get("batch no")
    if batch and not str(batch).upper().startswith("XXX"):
        notes_parts.append(f"Batch: {batch}")

    existing = db.query(CrmCustomer).filter(CrmCustomer.full_name == name[:150]).first()
    if existing:
        if address and not address.upper().startswith("XXX"):
            existing.address = address
        if notes_parts:
            existing.notes = ((existing.notes or "") + " | " + "; ".join(notes_parts)).strip(" |")
        existing.source = "coa"
        db.commit()
        return {
            "created": 0,
            "updated": 1,
            "skipped": 0,
            "total": 1,
            "errors": [],
            "customer_id": existing.id,
            "full_name": existing.full_name,
        }

    c = CrmCustomer(
        full_name=name[:150],
        address=None if not address or address.upper().startswith("XXX") else address,
        notes="; ".join(notes_parts) or None,
        source="coa",
        marketing_opt_in=False,
        is_active=True,
    )
    db.add(c)
    db.commit()
    db.refresh(c)
    return {
        "created": 1,
        "updated": 0,
        "skipped": 0,
        "total": 1,
        "errors": [],
        "customer_id": c.id,
        "full_name": c.full_name,
    }
